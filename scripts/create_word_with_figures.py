#!/usr/bin/env python3
"""
Create Word document with embedded figures at their legend locations.

Reads the markdown manuscript, parses it into structured content,
inserts each figure image immediately before its legend paragraph,
and outputs a formatted Word document.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


BASE = Path(__file__).resolve().parent.parent
MANUSCRIPT = BASE / "docs" / "manuscript" / "Chaco_Signaling_Manuscript.md"
FIGURES_DIR = BASE / "figures" / "final"
OUTPUT = BASE / "docs" / "manuscript" / "Chaco_Signaling_Manuscript.docx"

# Map figure numbers to filenames
FIGURE_FILES = {}
for f in sorted(FIGURES_DIR.glob("Figure_*.png")):
    m = re.search(r'Figure_(\d+)', f.name)
    if m:
        FIGURE_FILES[int(m.group(1))] = f


def setup_styles(doc):
    """Configure document styles: Times New Roman 11pt throughout."""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, name in enumerate(['Heading 1', 'Heading 2', 'Heading 3'], 1):
        h = doc.styles[name]
        h.font.name = 'Times New Roman'
        h.font.size = Pt(11)
        h.font.bold = (level == 1)
        h.font.italic = (level == 2)
        if level == 3:
            h.font.underline = True
            h.font.bold = False
            h.font.italic = False


def add_formatted_text(paragraph, text):
    """Add text with basic markdown bold/italic handling."""
    # Process bold+italic (***text***), bold (**text**), italic (*text*)
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+?\*|\$[^$]+\$)'
    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('$') and part.endswith('$'):
            # Math: just render as italic text (best we can do in docx)
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)


def is_figure_legend(text):
    """Check if a line is a figure legend (starts with **Figure N.**)."""
    return bool(re.match(r'\*\*Figure\s+\d+\.?\*?\*?', text.strip()))


def get_figure_num_from_legend(text):
    """Extract figure number from a legend line."""
    m = re.search(r'Figure\s+(\d+)', text)
    return int(m.group(1)) if m else None


def build_document():
    """Parse markdown and build Word document with embedded figures."""
    with open(MANUSCRIPT) as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = content.split('\n')
    para_buffer = []
    figures_inserted = set()

    def flush_paragraph():
        """Write accumulated paragraph text to the document."""
        if not para_buffer:
            return
        text = ' '.join(para_buffer)
        para_buffer.clear()

        if not text.strip():
            return

        # Check if this is a figure legend
        if is_figure_legend(text):
            fig_num = get_figure_num_from_legend(text)
            if fig_num and fig_num in FIGURE_FILES and fig_num not in figures_inserted:
                # Insert the figure image
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(str(FIGURE_FILES[fig_num]), width=Inches(6))
                figures_inserted.add(fig_num)

            # Now add the caption
            caption_para = doc.add_paragraph()
            add_formatted_text(caption_para, text)
            caption_para.paragraph_format.space_before = Pt(6)
            caption_para.paragraph_format.space_after = Pt(12)
            return

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, text)

    for line in lines:
        stripped = line.strip()

        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            flush_paragraph()
            p = doc.add_paragraph(stripped[2:], style='Heading 1')

        # Heading 2 (## )
        elif line.startswith('## '):
            flush_paragraph()
            p = doc.add_paragraph(stripped[3:], style='Heading 2')

        # Heading 3 (### )
        elif line.startswith('### '):
            flush_paragraph()
            p = doc.add_paragraph(stripped[4:], style='Heading 3')

        # Heading 4 (#### )
        elif line.startswith('#### '):
            flush_paragraph()
            p = doc.add_paragraph(stripped[5:], style='Heading 3')

        # Math display ($$...$$)
        elif stripped.startswith('$$'):
            flush_paragraph()
            math_text = stripped.strip('$').strip()
            if math_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(math_text)
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

        # Table rows (|...|)
        elif stripped.startswith('|'):
            # Skip markdown tables (they don't render well in docx via this method)
            # Could implement later with python-docx tables
            pass

        # Empty line
        elif stripped == '':
            flush_paragraph()

        # Regular text
        else:
            para_buffer.append(stripped)

    flush_paragraph()

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Figures embedded: {len(figures_inserted)} of {len(FIGURE_FILES)}")
    if figures_inserted:
        print(f"  Inserted: {sorted(figures_inserted)}")
    missing = set(FIGURE_FILES.keys()) - figures_inserted
    if missing:
        print(f"  Not inserted (no legend match): {sorted(missing)}")


if __name__ == '__main__':
    build_document()
